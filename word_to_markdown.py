#!/usr/bin/env python3
"""
Word to Markdown Converter

Converts Word documents (.docx) to Markdown format while preserving formatting
like bold, italics, and basic structure. Outputs with empty YAML front matter
headers that can be filled in.

Usage:
    python word_to_markdown.py -i input.docx -o output.md
    python word_to_markdown.py -i input.docx  # outputs to current directory
"""

import argparse
import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)


def extract_text_with_formatting(paragraph):
    """Extract text from a paragraph while preserving bold and italic formatting."""
    markdown_text = ""
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        markdown_text += text
    
    return markdown_text


def convert_paragraph_to_markdown(paragraph):
    """Convert a Word paragraph to Markdown format."""
    text = extract_text_with_formatting(paragraph).strip()
    
    if not text:
        return ""
    
    # Handle different paragraph styles
    style_name = paragraph.style.name.lower()
    
    # Headers
    if 'heading 1' in style_name or 'title' in style_name:
        return f"# {text}\n"
    elif 'heading 2' in style_name:
        return f"## {text}\n"
    elif 'heading 3' in style_name:
        return f"### {text}\n"
    elif 'heading 4' in style_name:
        return f"#### {text}\n"
    elif 'heading 5' in style_name:
        return f"##### {text}\n"
    elif 'heading 6' in style_name:
        return f"###### {text}\n"
    
    # Quote/Block quote
    elif 'quote' in style_name:
        return f"> {text}\n"
    
    # Regular paragraph
    else:
        return f"{text}\n"


def generate_yaml_frontmatter(title="", date=None, tags=None, excerpt=""):
    """Generate YAML front matter for the markdown file."""
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d %H:%M:%S %z")
    
    if tags is None:
        tags = "[]"
    elif isinstance(tags, list):
        tags = str(tags)
    
    frontmatter = f"""---
layout: post
title: "{title}"
date: {date}
tags: {tags}
excerpt: "{excerpt}"
---

"""
    return frontmatter


def convert_docx_to_markdown(input_path, output_path=None):
    """Convert a Word document to Markdown format."""
    try:
        # Load the Word document
        doc = Document(input_path)
        
        # Extract title from first paragraph or filename
        title = ""
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            title = doc.paragraphs[0].text.strip()
        else:
            title = Path(input_path).stem
        
        # Generate output filename if not provided
        if output_path is None:
            # Create filename from title, replacing spaces with underscores
            if title:
                # Clean the title for use as filename
                clean_filename = title.replace(' ', '_').replace('/', '_').replace('\\', '_')
                # Remove other problematic characters for filenames
                clean_filename = ''.join(c for c in clean_filename if c.isalnum() or c in ('_', '-', '.'))
                output_path = f"{clean_filename}.md"
            else:
                # Fallback to input filename if no title found
                output_path = Path(input_path).with_suffix('.md').name
        
        # Convert paragraphs to markdown
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            md_paragraph = convert_paragraph_to_markdown(paragraph)
            if md_paragraph:
                markdown_content.append(md_paragraph)
        
        # Join content and add spacing between paragraphs
        body_content = ""
        for i, content in enumerate(markdown_content):
            body_content += content
            # Add extra newline after paragraphs (but not headers)
            if not content.startswith('#') and i < len(markdown_content) - 1:
                body_content += "\n"
        
        # Generate complete markdown with front matter
        yaml_front = generate_yaml_frontmatter(title=title)
        full_markdown = yaml_front + body_content
        
        # Add footer similar to your example
        full_markdown += "\n---\n\n*[Add your footer text here]*\n"
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_markdown)
        
        print(f"Successfully converted '{input_path}' to '{output_path}'")
        return True
        
    except Exception as e:
        print(f"Error converting document: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description="Convert Word documents to Markdown format",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python word_to_markdown.py -i document.docx -o output.md
  python word_to_markdown.py -i document.docx
  python word_to_markdown.py -i /path/to/document.docx -o /path/to/output.md
        """
    )
    
    parser.add_argument(
        '-i', '--input',
        required=True,
        help='Path to input Word document (.docx)'
    )
    
    parser.add_argument(
        '-o', '--output',
        help='Path to output Markdown file (default: same name as input with .md extension in current directory)'
    )
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file '{args.input}' does not exist.")
        sys.exit(1)
    
    if not input_path.suffix.lower() in ['.docx', '.doc']:
        print(f"Warning: Input file '{args.input}' may not be a Word document.")
    
    # Determine output path
    output_path = args.output
    if output_path is None:
        output_path = input_path.with_suffix('.md').name
    
    # Convert the document
    success = convert_docx_to_markdown(str(input_path), output_path)
    
    if success:
        print(f"\nMarkdown file created: {output_path}")
        print("You can now edit the YAML front matter with your desired values.")
    else:
        sys.exit(1)


if __name__ == "__main__":
    main() 